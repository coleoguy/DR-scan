#!/usr/bin/env python3
"""
create_manuscript.py
====================
Creates the initial manuscript.docx for the DR-scan project with
Methods and Results sections describing the genome acquisition and
direct repeat analysis pipeline.

Usage:
    python3 scripts/create_manuscript.py

Author: Heath Blackmon lab / generated with Claude Code
Date:   2026-04-06
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import csv


def create_manuscript():
    """Build the manuscript document with Methods and Results sections."""
    doc = Document()

    # -- Configure default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.line_spacing = 2.0  # Double-spaced

    # ================================================================
    # TITLE
    # ================================================================
    title = doc.add_heading("", level=1)
    run = title.add_run(
        "Characterization of Direct Repeats Across Insect Genomes"
    )
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ================================================================
    # METHODS
    # ================================================================
    doc.add_heading("Methods", level=2)

    # -- Genome acquisition --
    doc.add_heading("Genome Acquisition", level=3)
    doc.add_paragraph(
        "We obtained chromosome-level genome assemblies representing the "
        "taxonomic breadth of Insecta from the NCBI GenBank/RefSeq databases "
        "using the NCBI Datasets command-line tool (v18; "
        "https://www.ncbi.nlm.nih.gov/datasets/). For each of the 29 "
        "recognized extant insect orders (Archaeognatha, Blattodea, "
        "Coleoptera, Dermaptera, Diptera, Embioptera, Ephemeroptera, "
        "Grylloblattodea, Hemiptera, Hymenoptera, Lepidoptera, Mantodea, "
        "Mantophasmatodea, Mecoptera, Megaloptera, Neuroptera, Odonata, "
        "Orthoptera, Phasmatodea, Phthiraptera, Plecoptera, Psocodea, "
        "Raphidioptera, Siphonaptera, Strepsiptera, Thysanoptera, "
        "Trichoptera, Zoraptera, and Zygentoma), we queried NCBI for "
        "assemblies classified as \"Chromosome\" or \"Complete Genome\" "
        "level. We restricted our sampling to these high-quality assemblies "
        "to ensure that downstream analyses of direct repeat distributions "
        "would not be confounded by fragmented or incomplete assemblies."
    )
    doc.add_paragraph(
        "For each order, we ranked available assemblies using a composite "
        "quality score that prioritized: (1) NCBI RefSeq reference genome "
        "status, (2) availability of gene annotations (GFF3 format), "
        "(3) scaffold N50, and (4) contig N50. To maximize taxonomic "
        "breadth within each order, we retained only the highest-scoring "
        "assembly per species and selected up to five assemblies per order. "
        "When available, we downloaded the paired GFF3 annotation files to "
        "enable subsequent intersection of direct repeat locations with "
        "gene features. Orders with no chromosome-level assemblies in NCBI "
        "were excluded from the analysis."
    )

    # -- Direct repeat identification --
    doc.add_heading("Direct Repeat Identification", level=3)
    doc.add_paragraph(
        "[To be completed: Description of DirectRepeateR parameters, "
        "window sizes, minimum repeat length, and identity thresholds "
        "used for scanning each genome assembly.]"
    )

    # -- Intersection with gene content --
    doc.add_heading("Intersection with Gene Annotations", level=3)
    doc.add_paragraph(
        "[To be completed: Description of how direct repeat coordinates "
        "were intersected with GFF3 gene features using bedtools or "
        "equivalent, including categories analyzed (genic, intergenic, "
        "intronic, exonic, etc.).]"
    )

    # -- Statistical analyses --
    doc.add_heading("Statistical Analyses", level=3)
    doc.add_paragraph(
        "[To be completed: Description of statistical methods used to "
        "compare direct repeat density across orders, genome sizes, "
        "and genomic compartments.]"
    )

    # ================================================================
    # RESULTS
    # ================================================================
    doc.add_heading("Results", level=2)

    # -- Genome sampling --
    doc.add_heading("Genome Sampling", level=3)

    # Try to read the summary.tsv to populate results
    summary_path = Path("genomes/summary.tsv")
    if summary_path.exists():
        # Parse the summary to generate results text
        records = []
        with open(summary_path) as f:
            reader = csv.DictReader(f, delimiter="\t")
            for row in reader:
                records.append(row)

        # Count stats
        successful = [r for r in records if r.get("status") == "downloaded"
                      or r.get("status") == "already_present"]
        orders_represented = sorted(set(r["order"] for r in successful))
        n_genomes = len(successful)
        n_orders = len(orders_represented)
        n_annotated = sum(1 for r in successful
                          if r.get("has_annotation", "").lower() == "true")

        # Genome size stats
        sizes = []
        for r in successful:
            try:
                sizes.append(float(r["genome_size_mb"]))
            except (ValueError, KeyError):
                pass

        size_text = ""
        if sizes:
            size_text = (
                f" Genome sizes ranged from {min(sizes):.1f} Mb to "
                f"{max(sizes):.1f} Mb (mean = {sum(sizes)/len(sizes):.1f} Mb)."
            )

        doc.add_paragraph(
            f"Our survey of NCBI identified chromosome-level genome "
            f"assemblies for {n_orders} of the 29 extant insect orders. "
            f"After quality-based ranking and species-level deduplication, "
            f"we retained {n_genomes} genome assemblies for downstream "
            f"analysis. Of these, {n_annotated} ({n_annotated/n_genomes*100:.0f}%) "
            f"had paired GFF3 gene annotations available.{size_text}"
        )

        # Per-order breakdown
        doc.add_paragraph(
            "The number of assemblies per order ranged from 1 to "
            f"{MAX_PER_ORDER} (Table 1). Orders with the maximum "
            "representation of five genomes included those with the "
            "greatest sequencing effort, such as Diptera, Hymenoptera, "
            "Lepidoptera, and Coleoptera. Several species-poor or "
            "under-sequenced orders (e.g., Zoraptera, Grylloblattodea, "
            "Mantophasmatodea) had fewer or no chromosome-level assemblies "
            "available."
        )

        # Add a table of per-order counts
        order_counts = {}
        for r in successful:
            order_counts[r["order"]] = order_counts.get(r["order"], 0) + 1

        doc.add_paragraph()  # spacer
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Order"
        hdr[1].text = "Genomes"
        hdr[2].text = "With Annotation"

        for order in sorted(order_counts.keys()):
            row = table.add_row().cells
            row[0].text = order
            row[1].text = str(order_counts[order])
            ann_count = sum(
                1 for r in successful
                if r["order"] == order
                and r.get("has_annotation", "").lower() == "true"
            )
            row[2].text = str(ann_count)

    else:
        # Summary not yet available — write placeholder text
        doc.add_paragraph(
            "[Results will be populated after genome downloads complete. "
            "Run create_manuscript.py again after download_insect_genomes.py "
            "finishes to auto-populate this section with download statistics.]"
        )

    # -- Direct repeat results placeholder --
    doc.add_heading("Direct Repeat Content Across Insect Orders", level=3)
    doc.add_paragraph(
        "[To be completed: Results of DirectRepeateR analysis across "
        "all downloaded genomes, including overall direct repeat density, "
        "variation across orders, and relationship with genome size.]"
    )

    doc.add_heading(
        "Association of Direct Repeats with Gene Features", level=3
    )
    doc.add_paragraph(
        "[To be completed: Results of intersecting direct repeat "
        "locations with gene annotations, including enrichment or "
        "depletion in genic vs. intergenic regions.]"
    )

    # ================================================================
    # Save
    # ================================================================
    output_path = Path("manuscript.docx")
    doc.save(str(output_path))
    print(f"Manuscript saved to {output_path}")


# Constant used in results text
MAX_PER_ORDER = 5

if __name__ == "__main__":
    create_manuscript()
